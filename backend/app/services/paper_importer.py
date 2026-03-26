from __future__ import annotations

from dataclasses import asdict, dataclass, field
import csv
import json
from pathlib import Path
import re
from typing import Any

from sqlalchemy.orm import Session

from app.core.exceptions import BadRequestException, NotFoundException
from app.models import Chapter, Paper, PaperQuestion, Question, Subject


OPTION_KEYS = ["A", "B", "C", "D", "E", "F", "G", "H"]
SECTION_TYPE_MAP = {
    "单选": "single_choice",
    "多选": "multiple_choice",
    "判断": "true_false",
    "简答": "short_answer",
    "操作": "short_answer",
    "填空": "short_answer",
    "综合": "short_answer",
}
TYPE_LABEL_MAP = {
    "single_choice": "单选题",
    "multiple_choice": "多选题",
    "true_false": "判断题",
    "short_answer": "简答/操作题",
}
TRUE_VALUES = {"T", "TRUE", "Y", "YES", "对", "正确", "√"}
FALSE_VALUES = {"F", "FALSE", "N", "NO", "错", "错误", "×"}


@dataclass
class NormalizedQuestion:
    sort_order: int
    question_type: str
    stem: str
    answer: list[str]
    explanation: str = ""
    options: list[dict[str, str]] | None = None
    difficulty: str = "medium"
    tags: list[str] = field(default_factory=list)
    score: int | None = None
    section_name: str | None = None
    chapter: str | None = None


@dataclass
class NormalizedPaper:
    subject_code: str
    code: str
    year: int
    season: str
    title: str
    description: str = ""
    duration_minutes: int = 90
    total_score: int | None = None
    source_type: str = "past_paper"
    questions: list[NormalizedQuestion] = field(default_factory=list)


@dataclass
class ImportReport:
    papers_created: int = 0
    papers_updated: int = 0
    questions_created: int = 0
    questions_updated: int = 0
    paper_questions_created: int = 0
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def default_score(question_type: str) -> int:
    if question_type == "multiple_choice":
        return 3
    if question_type == "short_answer":
        return 5
    return 2


def normalize_text(value: Any) -> str:
    return str(value or "").strip()


def normalize_tags(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [normalize_text(item) for item in value if normalize_text(item)]
    raw = normalize_text(value)
    if not raw:
        return []
    return [item.strip() for item in re.split(r"[、,，/；;|]", raw) if item.strip()]


def normalize_answers(question_type: str, value: Any) -> list[str]:
    if isinstance(value, list):
        items = [normalize_text(item) for item in value if normalize_text(item)]
    else:
        raw = normalize_text(value)
        if not raw:
            items = []
        elif question_type in {"single_choice", "multiple_choice"}:
            items = [item for item in re.split(r"[\s,，、/]+", raw.upper()) if item]
        elif question_type == "true_false":
            normalized = raw.upper()
            if normalized in TRUE_VALUES:
                items = ["T"]
            elif normalized in FALSE_VALUES:
                items = ["F"]
            else:
                items = [normalized]
        else:
            items = [raw]

    if question_type in {"single_choice", "multiple_choice", "true_false"}:
        return sorted(dict.fromkeys(item.upper() for item in items))
    return items


def infer_question_type(section_name: str | None, answer: Any, options: list[dict[str, str]] | None) -> str:
    if section_name:
        for key, value in SECTION_TYPE_MAP.items():
            if key in section_name:
                return value

    raw_answer = normalize_text(answer).upper()
    tokens = [item for item in re.split(r"[\s,，、/]+", raw_answer) if item]
    if raw_answer in TRUE_VALUES | FALSE_VALUES:
        return "true_false"
    if len(tokens) > 1 and all(token in OPTION_KEYS for token in tokens):
        return "multiple_choice"
    if len(tokens) == 1 and tokens[0] in OPTION_KEYS and options:
        return "single_choice"
    if options:
        return "single_choice"
    return "short_answer"


def build_options_from_row(row: dict[str, Any]) -> list[dict[str, str]] | None:
    options: list[dict[str, str]] = []
    for key in OPTION_KEYS:
        text = normalize_text(row.get(f"option_{key.lower()}") or row.get(f"option_{key}") or row.get(key))
        if text:
            options.append({"key": key, "text": text})
    return options or None


def build_paper_code(subject_code: str, year: int, season: str, title: str) -> str:
    season_token = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]", "", season.upper()) or "GEN"
    title_token = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "-", title.upper()).strip("-")[:32] or "PAPER"
    return f"{subject_code}-{year}-{season_token}-{title_token}"


def detect_format(path: Path) -> str:
    return {
        ".json": "json",
        ".xlsx": "xlsx",
        ".xlsm": "xlsx",
        ".csv": "csv",
        ".docx": "docx",
        ".pdf": "pdf",
        ".txt": "text",
    }.get(path.suffix.lower(), "unknown")


def parse_json_file(path: Path, subject_code_override: str | None = None) -> list[NormalizedPaper]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    papers_payload = payload["papers"] if isinstance(payload, dict) and "papers" in payload else [payload]
    papers: list[NormalizedPaper] = []

    for paper_payload in papers_payload:
        subject_code = subject_code_override or normalize_text(paper_payload.get("subject_code"))
        if not subject_code:
            raise BadRequestException("subject_code is required in JSON paper payload.")

        paper = NormalizedPaper(
            subject_code=subject_code,
            code=paper_payload.get("code")
            or build_paper_code(subject_code, int(paper_payload["year"]), normalize_text(paper_payload["season"]), normalize_text(paper_payload["title"])),
            year=int(paper_payload["year"]),
            season=normalize_text(paper_payload["season"]),
            title=normalize_text(paper_payload["title"]),
            description=normalize_text(paper_payload.get("description")),
            duration_minutes=int(paper_payload.get("duration_minutes") or 90),
            total_score=int(paper_payload["total_score"]) if normalize_text(paper_payload.get("total_score")) else None,
            source_type=normalize_text(paper_payload.get("source_type") or "past_paper"),
        )

        for index, question_payload in enumerate(paper_payload.get("questions", []), start=1):
            options = question_payload.get("options")
            if isinstance(options, dict):
                options = [{"key": key, "text": normalize_text(value)} for key, value in options.items() if normalize_text(value)]
            question_type = question_payload.get("question_type") or infer_question_type(
                question_payload.get("section_name"),
                question_payload.get("answer"),
                options,
            )
            paper.questions.append(
                NormalizedQuestion(
                    sort_order=int(question_payload.get("sort_order") or index),
                    question_type=question_type,
                    stem=normalize_text(question_payload.get("stem")),
                    answer=normalize_answers(question_type, question_payload.get("answer")),
                    explanation=normalize_text(question_payload.get("explanation")),
                    options=options,
                    difficulty=normalize_text(question_payload.get("difficulty") or "medium"),
                    tags=normalize_tags(question_payload.get("tags")),
                    score=int(question_payload["score"]) if normalize_text(question_payload.get("score")) else None,
                    section_name=normalize_text(question_payload.get("section_name")) or TYPE_LABEL_MAP.get(question_type),
                    chapter=normalize_text(question_payload.get("chapter")),
                )
            )
        papers.append(paper)
    return papers


def parse_excel_like_rows(rows: list[dict[str, Any]], paper_meta_map: dict[str, dict[str, Any]], subject_code_override: str | None = None) -> list[NormalizedPaper]:
    papers: dict[str, NormalizedPaper] = {}
    for row in rows:
        paper_code = normalize_text(row.get("paper_code"))
        if not paper_code:
            raise BadRequestException("Question rows require paper_code.")
        paper_meta = paper_meta_map.get(paper_code)
        if not paper_meta:
            raise BadRequestException(f"paper_code '{paper_code}' not found in paper metadata.")

        if paper_code not in papers:
            subject_code = subject_code_override or normalize_text(paper_meta.get("subject_code"))
            if not subject_code:
                raise BadRequestException(f"paper_code '{paper_code}' missing subject_code.")
            papers[paper_code] = NormalizedPaper(
                subject_code=subject_code,
                code=paper_code,
                year=int(paper_meta["year"]),
                season=normalize_text(paper_meta["season"]),
                title=normalize_text(paper_meta["title"]),
                description=normalize_text(paper_meta.get("description")),
                duration_minutes=int(paper_meta.get("duration_minutes") or 90),
                total_score=int(paper_meta["total_score"]) if normalize_text(paper_meta.get("total_score")) else None,
                source_type=normalize_text(paper_meta.get("source_type") or "past_paper"),
            )

        options = build_options_from_row(row)
        question_type = normalize_text(row.get("question_type")) or infer_question_type(row.get("section_name"), row.get("answer"), options)
        papers[paper_code].questions.append(
            NormalizedQuestion(
                sort_order=int(row.get("sort_order") or len(papers[paper_code].questions) + 1),
                question_type=question_type,
                stem=normalize_text(row.get("stem")),
                answer=normalize_answers(question_type, row.get("answer")),
                explanation=normalize_text(row.get("explanation")),
                options=options,
                difficulty=normalize_text(row.get("difficulty") or "medium"),
                tags=normalize_tags(row.get("tags")),
                score=int(row["score"]) if normalize_text(row.get("score")) else None,
                section_name=normalize_text(row.get("section_name")) or TYPE_LABEL_MAP.get(question_type),
                chapter=normalize_text(row.get("chapter")),
            )
        )
    return list(papers.values())


def parse_xlsx_file(path: Path, subject_code_override: str | None = None) -> list[NormalizedPaper]:
    try:
        from openpyxl import load_workbook
    except ModuleNotFoundError as exc:
        raise BadRequestException("openpyxl is required to import .xlsx files. Please install project requirements first.") from exc

    workbook = load_workbook(path, data_only=True)
    if "papers" not in workbook.sheetnames or "questions" not in workbook.sheetnames:
        raise BadRequestException("Excel workbook must contain 'papers' and 'questions' sheets.")

    def sheet_to_rows(sheet_name: str) -> list[dict[str, Any]]:
        sheet = workbook[sheet_name]
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            return []
        headers = [normalize_text(item) for item in rows[0]]
        result: list[dict[str, Any]] = []
        for values in rows[1:]:
            if not any(value is not None and normalize_text(value) for value in values):
                continue
            result.append({headers[index]: values[index] for index in range(min(len(headers), len(values)))})
        return result

    paper_rows = sheet_to_rows("papers")
    question_rows = sheet_to_rows("questions")
    meta_map = {normalize_text(row.get("paper_code")): row for row in paper_rows}
    return parse_excel_like_rows(question_rows, meta_map, subject_code_override=subject_code_override)


def parse_csv_file(path: Path, subject_code_override: str | None = None) -> list[NormalizedPaper]:
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        rows = list(csv.DictReader(file))
    if not rows:
        raise BadRequestException("CSV file is empty.")

    paper_code = normalize_text(rows[0].get("paper_code"))
    if not paper_code:
        raise BadRequestException("CSV requires a 'paper_code' column.")

    first = rows[0]
    meta_map = {
        paper_code: {
            "paper_code": paper_code,
            "subject_code": subject_code_override or normalize_text(first.get("subject_code")),
            "year": first.get("year"),
            "season": first.get("season"),
            "title": first.get("title"),
            "description": first.get("description"),
            "duration_minutes": first.get("duration_minutes"),
            "total_score": first.get("total_score"),
            "source_type": first.get("source_type"),
        }
    }
    return parse_excel_like_rows(rows, meta_map, subject_code_override=subject_code_override)


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise BadRequestException("python-docx is required to import .docx files. Please install project requirements first.") from exc

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text))
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise BadRequestException("pypdf is required to import .pdf files. Please install project requirements first.") from exc

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def parse_plain_text_document(
    text: str,
    subject_code: str,
    paper_code: str | None = None,
    year: int | None = None,
    season: str | None = None,
    title: str | None = None,
) -> list[NormalizedPaper]:
    lines = [normalize_text(line) for line in text.splitlines()]
    lines = [line for line in lines if line]
    if not lines:
        raise BadRequestException("Document text is empty after extraction.")

    inferred_title = title or lines[0]
    inferred_year = year or next(
        (int(match.group(1)) for line in lines[:20] for match in [re.search(r"(20\d{2})", line)] if match),
        2024,
    )
    inferred_season = season or next(
        (match.group(1) for line in lines[:20] for match in [re.search(r"([39]月)", line)] if match),
        "9月",
    )

    paper = NormalizedPaper(
        subject_code=subject_code,
        code=paper_code or build_paper_code(subject_code, inferred_year, inferred_season, inferred_title),
        year=inferred_year,
        season=inferred_season,
        title=inferred_title,
        description="由文档自动解析生成，建议导入后人工抽样复核。",
        duration_minutes=90,
        source_type="past_paper",
    )

    current_section: str | None = None
    current_question: dict[str, Any] | None = None
    current_explanation_lines: list[str] = []
    option_pattern = re.compile(r"^\s*([A-H])[\.．、\s]+(.+)$")
    question_pattern = re.compile(r"^\s*(\d+)[\.．、\)]\s*(.+)$")

    def flush_question() -> None:
        nonlocal current_question, current_explanation_lines
        if not current_question:
            return
        if current_explanation_lines and not current_question.get("explanation"):
            current_question["explanation"] = " ".join(current_explanation_lines).strip()
        question_type = current_question.get("question_type") or infer_question_type(
            current_section,
            current_question.get("answer"),
            current_question.get("options"),
        )
        paper.questions.append(
            NormalizedQuestion(
                sort_order=int(current_question["sort_order"]),
                question_type=question_type,
                stem=normalize_text(current_question["stem"]),
                answer=normalize_answers(question_type, current_question.get("answer")),
                explanation=normalize_text(current_question.get("explanation")),
                options=current_question.get("options") or None,
                difficulty=normalize_text(current_question.get("difficulty") or "medium"),
                tags=normalize_tags(current_question.get("tags")),
                score=int(current_question["score"]) if current_question.get("score") else None,
                section_name=current_section or TYPE_LABEL_MAP.get(question_type),
                chapter=normalize_text(current_question.get("chapter")),
            )
        )
        current_question = None
        current_explanation_lines = []

    for line in lines:
        normalized_line = re.sub(r"\s+", " ", line)
        if any(keyword in normalized_line for keyword in SECTION_TYPE_MAP):
            flush_question()
            current_section = normalized_line
            continue

        question_match = question_pattern.match(normalized_line)
        if question_match:
            flush_question()
            current_question = {
                "sort_order": int(question_match.group(1)),
                "stem": question_match.group(2),
                "options": [],
            }
            continue

        if current_question is None:
            continue

        option_match = option_pattern.match(normalized_line)
        if option_match:
            current_question.setdefault("options", []).append(
                {"key": option_match.group(1), "text": option_match.group(2).strip()}
            )
            continue

        if normalized_line.startswith("答案") or normalized_line.startswith("参考答案"):
            answer_text = re.sub(r"^参考?答案[:：]?", "", normalized_line).strip()
            current_question["answer"] = answer_text
            continue

        if normalized_line.startswith("解析") or normalized_line.startswith("参考解析"):
            explanation_text = re.sub(r"^参考?解析[:：]?", "", normalized_line).strip()
            current_question["explanation"] = explanation_text
            continue

        if normalized_line.startswith("难度"):
            current_question["difficulty"] = re.sub(r"^难度[:：]?", "", normalized_line).strip() or "medium"
            continue

        if normalized_line.startswith("标签"):
            current_question["tags"] = normalize_tags(re.sub(r"^标签[:：]?", "", normalized_line).strip())
            continue

        if current_question.get("explanation"):
            current_question["explanation"] = f"{current_question['explanation']} {normalized_line}".strip()
        elif current_question.get("options"):
            current_explanation_lines.append(normalized_line)
        else:
            current_question["stem"] = f"{current_question['stem']} {normalized_line}".strip()

    flush_question()

    if not paper.questions:
        raise BadRequestException("No questions were recognized from document text. Please prefer JSON/Excel import or provide explicit metadata.")

    return [paper]


def parse_input_file(
    path: str | Path,
    *,
    input_format: str = "auto",
    subject_code: str | None = None,
    paper_code: str | None = None,
    year: int | None = None,
    season: str | None = None,
    title: str | None = None,
) -> list[NormalizedPaper]:
    file_path = Path(path)
    if not file_path.exists():
        raise BadRequestException(f"Input file not found: {file_path}")

    detected_format = detect_format(file_path) if input_format == "auto" else input_format

    if detected_format == "json":
        return parse_json_file(file_path, subject_code_override=subject_code)
    if detected_format == "xlsx":
        return parse_xlsx_file(file_path, subject_code_override=subject_code)
    if detected_format == "csv":
        return parse_csv_file(file_path, subject_code_override=subject_code)
    if detected_format == "docx":
        if not subject_code:
            raise BadRequestException("subject_code is required when importing .docx files.")
        text = extract_docx_text(file_path)
        return parse_plain_text_document(text, subject_code, paper_code, year, season, title)
    if detected_format == "pdf":
        if not subject_code:
            raise BadRequestException("subject_code is required when importing .pdf files.")
        text = extract_pdf_text(file_path)
        return parse_plain_text_document(text, subject_code, paper_code, year, season, title)
    if detected_format == "text":
        if not subject_code:
            raise BadRequestException("subject_code is required when importing text files.")
        text = file_path.read_text(encoding="utf-8")
        return parse_plain_text_document(text, subject_code, paper_code, year, season, title)

    raise BadRequestException(f"Unsupported input format: {detected_format}")


def serialize_normalized_papers(papers: list[NormalizedPaper]) -> dict[str, Any]:
    return {
        "papers": [
            {
                "subject_code": paper.subject_code,
                "code": paper.code,
                "year": paper.year,
                "season": paper.season,
                "title": paper.title,
                "description": paper.description,
                "duration_minutes": paper.duration_minutes,
                "total_score": paper.total_score,
                "source_type": paper.source_type,
                "questions": [
                    {
                        "sort_order": question.sort_order,
                        "question_type": question.question_type,
                        "section_name": question.section_name,
                        "chapter": question.chapter,
                        "stem": question.stem,
                        "options": question.options,
                        "answer": question.answer,
                        "explanation": question.explanation,
                        "difficulty": question.difficulty,
                        "tags": question.tags,
                        "score": question.score,
                    }
                    for question in paper.questions
                ],
            }
            for paper in papers
        ]
    }


def get_or_create_import_chapter(db: Session, subject_id: int, chapter_title: str) -> Chapter:
    chapter = (
        db.query(Chapter)
        .filter(Chapter.subject_id == subject_id, Chapter.title == chapter_title)
        .one_or_none()
    )
    if chapter is not None:
        return chapter

    max_sort_order = db.query(Chapter).filter(Chapter.subject_id == subject_id).count() + 1
    chapter = Chapter(
        subject_id=subject_id,
        title=chapter_title,
        code=re.sub(r"[^A-Za-z0-9]+", "-", chapter_title.upper()).strip("-")[:48] or f"IMPORT-{subject_id}-{max_sort_order}",
        outline=f"{chapter_title}（由导入器自动创建）",
        sort_order=max_sort_order,
        estimated_minutes=45,
    )
    db.add(chapter)
    db.flush()
    return chapter


def get_subject_by_code_or_name(db: Session, subject_identifier: str) -> Subject | None:
    return (
        db.query(Subject)
        .filter((Subject.code == subject_identifier) | (Subject.name == subject_identifier))
        .one_or_none()
    )


def import_normalized_papers(
    db: Session,
    papers: list[NormalizedPaper],
    *,
    replace_paper_questions: bool = True,
) -> ImportReport:
    report = ImportReport()

    for normalized_paper in papers:
        subject = get_subject_by_code_or_name(db, normalized_paper.subject_code)
        if subject is None:
            raise NotFoundException(f"Subject '{normalized_paper.subject_code}' not found.")

        paper = db.query(Paper).filter(Paper.code == normalized_paper.code).one_or_none()
        if paper is None:
            paper = Paper(
                subject_id=subject.id,
                code=normalized_paper.code,
                year=normalized_paper.year,
                season=normalized_paper.season,
                title=normalized_paper.title,
                source_type=normalized_paper.source_type,
                description=normalized_paper.description,
                total_questions=0,
                total_score=normalized_paper.total_score or 0,
                duration_minutes=normalized_paper.duration_minutes,
            )
            db.add(paper)
            db.flush()
            report.papers_created += 1
        else:
            paper.subject_id = subject.id
            paper.year = normalized_paper.year
            paper.season = normalized_paper.season
            paper.title = normalized_paper.title
            paper.source_type = normalized_paper.source_type
            paper.description = normalized_paper.description
            paper.duration_minutes = normalized_paper.duration_minutes
            report.papers_updated += 1

        if replace_paper_questions:
            db.query(PaperQuestion).filter(PaperQuestion.paper_id == paper.id).delete()

        paper_total_score = 0

        for normalized_question in sorted(normalized_paper.questions, key=lambda item: item.sort_order):
            if not normalized_question.stem:
                report.warnings.append(f"{paper.code} 第 {normalized_question.sort_order} 题题干为空，已跳过。")
                continue

            chapter_title = (
                normalized_question.chapter
                or normalized_question.section_name
                or f"{normalized_paper.year}年{normalized_paper.season}真题"
            )
            chapter = get_or_create_import_chapter(db, subject.id, chapter_title)

            existing_question = (
                db.query(Question)
                .filter(
                    Question.subject_id == subject.id,
                    Question.question_type == normalized_question.question_type,
                    Question.stem == normalized_question.stem,
                )
                .one_or_none()
            )

            question_score = normalized_question.score or default_score(normalized_question.question_type)
            if existing_question is None:
                existing_question = Question(
                    subject_id=subject.id,
                    chapter_id=chapter.id,
                    chapter=chapter.title,
                    question_type=normalized_question.question_type,
                    stem=normalized_question.stem,
                    options=normalized_question.options,
                    answer=normalized_question.answer,
                    explanation=normalized_question.explanation,
                    difficulty=normalized_question.difficulty,
                    tags=normalized_question.tags,
                    score=question_score,
                )
                db.add(existing_question)
                db.flush()
                report.questions_created += 1
            else:
                existing_question.chapter_id = chapter.id
                existing_question.chapter = chapter.title
                existing_question.options = normalized_question.options
                existing_question.answer = normalized_question.answer
                existing_question.explanation = normalized_question.explanation
                existing_question.difficulty = normalized_question.difficulty
                existing_question.tags = normalized_question.tags
                existing_question.score = question_score
                report.questions_updated += 1

            paper_question = PaperQuestion(
                paper_id=paper.id,
                question_id=existing_question.id,
                sort_order=normalized_question.sort_order,
                section_name=normalized_question.section_name or TYPE_LABEL_MAP.get(normalized_question.question_type),
                score_override=question_score,
                note=None,
            )
            db.add(paper_question)
            report.paper_questions_created += 1
            paper_total_score += question_score

        paper.total_questions = len(normalized_paper.questions)
        paper.total_score = normalized_paper.total_score or paper_total_score

    db.commit()
    return report
